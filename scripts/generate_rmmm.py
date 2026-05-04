"""
RMMM (Risk Mitigation, Monitoring, and Management) Plan Generator
Project: Integrated Business Management Suite (IBMS)
Date: April 24, 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────
# RISK DATA
# ─────────────────────────────────────────────────

UNSORTED_RISKS = [
    {
        "id": 1,
        "risk": "Rapid scope expansion due to evolving enterprise requirements",
        "category": "BU",
        "probability": 70,
        "impact": 3,
        "rmmm": (
            "Define and freeze scope clearly; get stakeholder sign-off; "
            "use change-control board; monitor scope changes strictly; "
            "reassess budget/timeline for every approved change"
        ),
    },
    {
        "id": 2,
        "risk": "AI/ML service integration complexity causes schedule delays",
        "category": "TE",
        "probability": 65,
        "impact": 2,
        "rmmm": (
            "Spike AI/ML modules early; use containerised ML pipelines; "
            "allocate 20 % schedule buffer for model tuning; "
            "maintain fallback rule-based logic; review weekly"
        ),
    },
    {
        "id": 3,
        "risk": "Data security breach or unauthorised access to sensitive business data",
        "category": "PS",
        "probability": 55,
        "impact": 1,
        "rmmm": (
            "Enforce JWT + TOTP 2FA; implement RBAC with least-privilege; "
            "run quarterly penetration tests; apply Supabase RLS policies; "
            "maintain audit logs; conduct security training for team"
        ),
    },
    {
        "id": 4,
        "risk": "Third-party service outage (Supabase, Redis, Groq AI) disrupts operations",
        "category": "DE",
        "probability": 50,
        "impact": 2,
        "rmmm": (
            "Implement in-memory fallback for all ERP ops; use Redis caching; "
            "configure health-check endpoints; document manual override procedures; "
            "evaluate multi-provider redundancy"
        ),
    },
    {
        "id": 5,
        "risk": "Team skill gaps with advanced tech stack (FastAPI, Frappe, K8s, ML)",
        "category": "ST",
        "probability": 45,
        "impact": 2,
        "rmmm": (
            "Conduct training sessions on key technologies; pair experienced devs "
            "with juniors; use comprehensive documentation; allocate ramp-up buffer; "
            "consider specialist contractors for critical modules"
        ),
    },
    {
        "id": 6,
        "risk": "Performance degradation under high concurrent user load",
        "category": "PS",
        "probability": 40,
        "impact": 2,
        "rmmm": (
            "Load test with target concurrency before each release; "
            "use K8s HPA (3-replica default); optimise Supabase queries; "
            "tune Redis cache TTLs; profile WebSocket KPI broadcast bottlenecks"
        ),
    },
    {
        "id": 7,
        "risk": "Customer resistance to adopting the new integrated ERP platform",
        "category": "CU",
        "probability": 35,
        "impact": 3,
        "rmmm": (
            "Involve key stakeholders in design reviews; provide onboarding tutorials; "
            "offer phased rollout with parallel-run period; gather feedback early; "
            "assign a dedicated customer success manager"
        ),
    },
    {
        "id": 8,
        "risk": "Budget overrun due to underestimated infrastructure and AI API costs",
        "category": "BU",
        "probability": 30,
        "impact": 2,
        "rmmm": (
            "Perform detailed cloud-cost modelling before sprint starts; "
            "set budget alerts on GCP/AWS billing; review costs bi-weekly; "
            "include 15 % contingency reserve; explore open-source AI model alternatives"
        ),
    },
]

# Sort descending by probability
SORTED_RISKS = sorted(UNSORTED_RISKS, key=lambda r: r["probability"], reverse=True)

IMPACT_LABELS = {1: "Catastrophic", 2: "Critical", 3: "Marginal", 4: "Negligible"}
PROB_LABELS = {
    (80, 100): "Very High",
    (60, 79): "High",
    (40, 59): "Medium",
    (20, 39): "Low",
    (0, 19): "Very Low",
}

CATEGORY_FULL = {
    "PS": "Product Size",
    "BU": "Business Impact",
    "CU": "Customer Characteristics",
    "PD": "Process Definition",
    "DE": "Development Environment",
    "TE": "Technology to be Built",
    "ST": "Staff Size & Experience",
}


def prob_label(p):
    for (lo, hi), label in PROB_LABELS.items():
        if lo <= p <= hi:
            return label
    return ""


# ─────────────────────────────────────────────────
# COLOUR PALETTE
# ─────────────────────────────────────────────────

DARK_BLUE = RGBColor(0x1F, 0x3B, 0x6B)
MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLU = RGBColor(0xBD, 0xD7, 0xEE)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
LIGHT_GRN = RGBColor(0xE2, 0xEF, 0xDA)

# A4 page width = 21 cm; margins 2.54 cm each side → text area = 15.92 cm
PAGE_W = Cm(15.92)

# ─────────────────────────────────────────────────
# LOW-LEVEL XML HELPERS
# ─────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(rgb))
    tcPr.append(shd)


def set_table_fullwidth(tbl):
    """Force table to span 100 % of the text column."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), "5000")   # 5000 / 50 = 100 %
    tblPr.append(tblW)


def keep_table_on_page(tbl):
    """Ask Word to keep each row together (avoid mid-row page breaks)."""
    for row in tbl.rows:
        trPr = row._tr.get_or_add_trPr()
        cant = OxmlElement("w:cantSplit")
        cant.set(qn("w:val"), "1")
        trPr.append(cant)


def no_space_after(para):
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


# ─────────────────────────────────────────────────
# CELL / PARAGRAPH BUILDERS
# ─────────────────────────────────────────────────

def cell_para(cell, text, bold=False, italic=False, size=9.5,
              color: RGBColor = BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=1, space_after=1):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size       = Pt(size)
    run.font.color.rgb  = color
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(13)
    run.font.color.rgb = DARK_BLUE


def add_sub_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(10.5)
    run.font.color.rgb = MID_BLUE


def add_body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.italic         = italic
    run.font.size      = Pt(10)
    run.font.color.rgb = BLACK


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(9.5)


def set_tbl_style(tbl):
    tbl.style     = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fullwidth(tbl)

# ─────────────────────────────────────────────────
# DOCUMENT BUILDER
# ─────────────────────────────────────────────────

def build_document():
    doc = Document()

    # Standard 1-inch (2.54 cm) margins on all sides
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)

    # ── TITLE PAGE ──────────────────────────────────────────────────────────
    # Push content to vertical centre with top space
    for _ in range(6):
        sp = doc.add_paragraph()
        no_space_after(sp)

    def centered(text, size, bold=False, italic=False, color=DARK_BLUE, space_b=4, space_a=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_b)
        p.paragraph_format.space_after  = Pt(space_a)
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size)
        r.font.color.rgb = color

    centered("RMMM Plan", 30, bold=True, space_b=0, space_a=6)
    centered("Risk Mitigation, Monitoring, and Management Plan",
             15, italic=True, color=MID_BLUE, space_b=0, space_a=18)
    centered("Project: Integrated Business Management Suite (IBMS)",
             12, bold=True, color=BLACK, space_b=0, space_a=10)

    # Divider line via table
    div = doc.add_table(rows=1, cols=1)
    div.style = "Table Grid"
    set_table_fullwidth(div)
    dc = div.rows[0].cells[0]
    set_cell_bg(dc, MID_BLUE)
    cell_para(dc, "", size=2)

    # Team box
    team_tbl = doc.add_table(rows=6, cols=2)
    set_tbl_style(team_tbl)
    team_tbl.columns[0].width = Cm(4.5)
    team_tbl.columns[1].width = Cm(11.42)
    team_rows = [
        ("Version",  "1.0"),
        ("Date",     "April 24, 2026"),
        ("Course",   "SEPM (Software Engineering and Project Management)"),
        ("Team",     "Shivansh Srivastava   –   Product Developer\n"
                     "Prahallad Padhan       –   Product Owner\n"
                     "Ranveer Rai Khare      –   Scrum Master"),
        ("Project",  "Integrated Business Management Suite (IBMS)"),
        ("Status",   "Version 1.0  |  April 2026"),
    ]
    for ri, (lbl, val) in enumerate(team_rows):
        lc = team_tbl.rows[ri].cells[0]
        vc = team_tbl.rows[ri].cells[1]
        set_cell_bg(lc, DARK_BLUE)
        cell_para(lc, lbl, bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        bg = LIGHT_BLU if ri % 2 == 0 else WHITE
        set_cell_bg(vc, bg)
        cell_para(vc, val, size=10)

    doc.add_page_break()

    # ── PAGE 2 — SECTION 1: INTRODUCTION + REFERENCE TABLES ─────────────────
    add_heading(doc, "1.  Introduction")
    add_body(doc,
        "This document presents the Risk Mitigation, Monitoring, and Management (RMMM) Plan "
        "for the Integrated Business Management Suite (IBMS) — an AI-first Enterprise Resource "
        "Planning platform built on FastAPI, Frappe, Supabase, and machine-learning services.  "
        "The plan identifies key risks, assigns probability and impact values, and specifies "
        "concrete mitigation strategies for each risk.")
    add_body(doc, "The plan follows the three-step process (Pressman, 2015):")
    add_bullet(doc, "Step 1 – Form a Risk Table listing all identified risks.")
    add_bullet(doc, "Step 2 – Sort risks in descending order of probability.")
    add_bullet(doc, "Step 3 – Produce a Risk Information Sheet (RIS) for each individual risk.")

    # Risk Category Table
    add_sub_heading(doc, "Risk Category Reference")
    cat_data = [
        ("PS", "Product Size",            "Risks related to overall size, scale, and complexity of the software."),
        ("BU", "Business Impact",         "Risks from business constraints: market conditions, competition, management."),
        ("CU", "Customer Characteristics","Risks related to customer behaviour, expectations, and communication."),
        ("PD", "Process Definition",      "Risks from unclear, undefined, or poorly followed development processes."),
        ("DE", "Development Environment", "Risks related to availability, reliability, and quality of dev tools."),
        ("TE", "Technology to be Built",  "Risks from complexity, novelty, or uncertainty of the technology used."),
        ("ST", "Staff Size & Experience", "Risks related to skills, experience, and availability of the team."),
    ]
    cat_tbl = doc.add_table(rows=len(cat_data) + 1, cols=3)
    set_tbl_style(cat_tbl)
    cat_tbl.columns[0].width = Cm(1.8)
    cat_tbl.columns[1].width = Cm(4.3)
    cat_tbl.columns[2].width = Cm(9.82)
    for i, h in enumerate(["Short Form", "Risk Category", "Explanation"]):
        c = cat_tbl.rows[0].cells[i]
        set_cell_bg(c, MID_BLUE)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    for ri, (short, name, expl) in enumerate(cat_data, 1):
        bg = LIGHT_BLU if ri % 2 == 0 else WHITE
        cells = cat_tbl.rows[ri].cells
        set_cell_bg(cells[0], bg); cell_para(cells[0], short, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[1], bg); cell_para(cells[1], name, bold=True, size=9)
        set_cell_bg(cells[2], bg); cell_para(cells[2], expl, size=9)

    # Probability + Impact side-by-side using a 1-row wrapper table
    add_sub_heading(doc, "Probability Scale  &  Impact Severity")
    side = doc.add_table(rows=1, cols=2)
    set_table_fullwidth(side)
    side.columns[0].width = Cm(7.5)
    side.columns[1].width = Cm(8.42)

    # Left cell: probability table
    lc_outer = side.rows[0].cells[0]
    lc_outer.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lc_outer.paragraphs[0].clear()
    prob_tbl = lc_outer.add_table(rows=6, cols=2)
    prob_tbl.style = "Table Grid"
    prob_rows_data = [
        ("Range", "Interpretation", True),
        ("80% – 100%", "Very High", False),
        ("60% – 79%",  "High",      False),
        ("40% – 59%",  "Medium",    False),
        ("20% – 39%",  "Low",       False),
        ("< 20%",      "Very Low",  False),
    ]
    prob_tbl.columns[0].width = Cm(3.5)
    prob_tbl.columns[1].width = Cm(4.0)
    for ri, (a, b, is_hdr) in enumerate(prob_rows_data):
        ca = prob_tbl.rows[ri].cells[0]
        cb = prob_tbl.rows[ri].cells[1]
        bg_hdr = MID_BLUE if is_hdr else (LIGHT_BLU if ri % 2 == 0 else WHITE)
        set_cell_bg(ca, bg_hdr); cell_para(ca, a, bold=is_hdr, color=WHITE if is_hdr else BLACK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cb, bg_hdr); cell_para(cb, b, bold=is_hdr, color=WHITE if is_hdr else BLACK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Right cell: impact table
    rc_outer = side.rows[0].cells[1]
    rc_outer.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rc_outer.paragraphs[0].clear()
    imp_tbl = rc_outer.add_table(rows=5, cols=3)
    imp_tbl.style = "Table Grid"
    imp_rows_data = [
        ("Lvl", "Label",         "Description",                              True),
        ("1",   "Catastrophic",  "Mission failure; financial loss > $500K",  False),
        ("2",   "Critical",      "Significant degradation; $100K – $500K",   False),
        ("3",   "Marginal",      "Minor degradation; $1K – $100K",           False),
        ("4",   "Negligible",    "Minor inconvenience; < $1K",               False),
    ]
    imp_tbl.columns[0].width = Cm(1.0)
    imp_tbl.columns[1].width = Cm(3.0)
    imp_tbl.columns[2].width = Cm(4.42)
    for ri, (lvl, lbl, desc, is_hdr) in enumerate(imp_rows_data):
        cl, cn, cd = [imp_tbl.rows[ri].cells[k] for k in range(3)]
        bg_hdr = MID_BLUE if is_hdr else (LIGHT_BLU if ri % 2 == 0 else WHITE)
        set_cell_bg(cl, bg_hdr); cell_para(cl, lvl,  bold=is_hdr, color=WHITE if is_hdr else BLACK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cn, bg_hdr); cell_para(cn, lbl,  bold=is_hdr, color=WHITE if is_hdr else BLACK, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cd, bg_hdr); cell_para(cd, desc, bold=is_hdr, color=WHITE if is_hdr else BLACK, size=9)

    doc.add_page_break()

    # ── PAGE 3 — STEP 1: RISK TABLE (UNSORTED) ───────────────────────────────
    add_heading(doc, "2.  Step 1 – Risk Table (Initial Identification)")
    add_body(doc,
        "All identified risks are listed below with a unique ID, description, category, "
        "estimated probability, impact level, and a brief RMMM reference.")

    risk_tbl = doc.add_table(rows=len(UNSORTED_RISKS) + 1, cols=6)
    set_tbl_style(risk_tbl)
    # Total = 15.92 cm
    widths = [Cm(1.1), Cm(5.3), Cm(1.6), Cm(1.9), Cm(1.9), Cm(4.12)]
    for i, w in enumerate(widths):
        risk_tbl.columns[i].width = w

    col_hdrs = ["Risk\nID", "Risk Description", "Cat.", "Probability", "Impact", "RMMM Summary"]
    for i, h in enumerate(col_hdrs):
        c = risk_tbl.rows[0].cells[i]
        set_cell_bg(c, DARK_BLUE)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, risk in enumerate(UNSORTED_RISKS, 1):
        bg = LIGHT_BLU if ri % 2 == 0 else WHITE
        cells = risk_tbl.rows[ri].cells
        set_cell_bg(cells[0], bg); cell_para(cells[0], f"R{risk['id']}", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[1], bg); cell_para(cells[1], risk["risk"], size=9)
        set_cell_bg(cells[2], bg); cell_para(cells[2], risk["category"], bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[3], bg); cell_para(cells[3], f"{risk['probability']}%\n({prob_label(risk['probability'])})", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[4], bg); cell_para(cells[4], f"{risk['impact']} – {IMPACT_LABELS[risk['impact']]}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        # Truncate RMMM summary to ~80 chars so column stays narrow
        summary = risk["rmmm"]
        if len(summary) > 78:
            summary = summary[:75].rstrip() + "…"
        set_cell_bg(cells[5], bg); cell_para(cells[5], summary, size=8.5)

    keep_table_on_page(risk_tbl)
    doc.add_page_break()

    # ── PAGE 4 — STEP 2: SORTED RISK TABLE ───────────────────────────────────
    add_heading(doc, "3.  Step 2 – Risk Table Sorted by Probability (Descending)")
    add_body(doc,
        "Risks re-ordered in descending probability so that the highest-likelihood risks "
        "receive the most immediate management attention.")

    sorted_tbl = doc.add_table(rows=len(SORTED_RISKS) + 1, cols=5)
    set_tbl_style(sorted_tbl)
    # Total = 15.92 cm
    s_widths = [Cm(1.1), Cm(6.3), Cm(1.6), Cm(2.2), Cm(4.72)]
    for i, w in enumerate(s_widths):
        sorted_tbl.columns[i].width = w

    s_hdrs = ["Risk\nID", "Risk Description", "Cat.", "Probability", "Impact Level"]
    for i, h in enumerate(s_hdrs):
        c = sorted_tbl.rows[0].cells[i]
        set_cell_bg(c, DARK_BLUE)
        cell_para(c, h, bold=True, color=WHITE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, risk in enumerate(SORTED_RISKS, 1):
        bg = LIGHT_BLU if ri % 2 == 0 else WHITE
        cells = sorted_tbl.rows[ri].cells
        set_cell_bg(cells[0], bg); cell_para(cells[0], f"R{risk['id']}", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[1], bg); cell_para(cells[1], risk["risk"], size=9)
        set_cell_bg(cells[2], bg); cell_para(cells[2], risk["category"], bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[3], bg); cell_para(cells[3], f"{risk['probability']}%\n({prob_label(risk['probability'])})", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(cells[4], bg); cell_para(cells[4], f"{risk['impact']} – {IMPACT_LABELS[risk['impact']]}", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    keep_table_on_page(sorted_tbl)
    doc.add_page_break()

    # ── PAGES 5-12 — STEP 3: RISK INFORMATION SHEETS ─────────────────────────
    add_heading(doc, "4.  Step 3 – Risk Information Sheets (RIS)")
    add_body(doc,
        "A Risk Information Sheet is produced for each of the eight identified risks.  "
        "Each sheet details sub-conditions, mitigation strategies, contingency plans, "
        "trigger criteria, and ownership assignments.")

    doc.add_page_break()

    RIS_DATA = [
        {
            "id": 1,
            "category": "BU",
            "prob": 70,
            "impact": 3,
            "description": (
                "Stakeholders continuously request additional features beyond the agreed scope, "
                "resulting in schedule slippage and cost overrun."
            ),
            "subconditions": [
                "Requirements were collected informally without formal sign-off.",
                "Multiple stakeholder groups have conflicting priorities.",
                "Market changes prompt leadership to request mid-sprint additions.",
            ],
            "mitigation": [
                "Establish a formal Change Control Board (CCB) with documented approval workflows.",
                "Baseline scope in a signed Software Requirements Specification (SRS).",
                "Use backlog grooming to assess impact of every change before acceptance.",
                "Set a sprint change-freeze policy: no new requirements after sprint planning.",
            ],
            "contingency": [
                "Renegotiate project timeline and budget for every approved change.",
                "Defer non-critical features to a future release roadmap.",
                "Escalate to project sponsor if scope growth exceeds 15 % of baseline.",
                "Initiate a formal re-baselining with an updated WBS.",
            ],
            "trigger": "Approved change requests exceed 10 % of baseline story points in any sprint.",
            "assigned_to": "Prahallad Padhan  (Product Owner)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
        {
            "id": 2,
            "category": "TE",
            "prob": 65,
            "impact": 2,
            "description": (
                "Integration of AI/ML services (anomaly detection, fraud scoring, dynamic pricing, "
                "predictive inventory, lead scoring) is more complex than estimated, causing schedule "
                "delays and performance issues."
            ),
            "subconditions": [
                "AI/ML libraries require dependency resolution across Python versions.",
                "Model training data quality is lower than expected.",
                "Real-time inference latency does not meet the ≤ 200 ms SLA.",
            ],
            "mitigation": [
                "Allocate a dedicated 2-week AI/ML spike at the start of each milestone.",
                "Containerise all ML services independently to isolate failures.",
                "Implement rule-based fallback logic for each ML service if a model is unavailable.",
                "Track model accuracy and latency metrics via the monitoring dashboard.",
            ],
            "contingency": [
                "Substitute unavailable ML features with enhanced rule-based approximations.",
                "Extend sprint timeline by up to 20 % buffer for AI integration blockers.",
                "Reduce AI feature scope to core functions only for the initial release.",
                "Engage an external ML consultant if internal team cannot resolve blockers.",
            ],
            "trigger": "AI/ML integration tasks remain blocked for more than 3 consecutive days.",
            "assigned_to": "Shivansh Srivastava  (Product Developer)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
        {
            "id": 3,
            "category": "PS",
            "prob": 55,
            "impact": 1,
            "description": (
                "Unauthorised access to sensitive business data (financial records, employee data, "
                "customer PII) due to exploited security vulnerabilities in the IBMS platform."
            ),
            "subconditions": [
                "JWT secret rotation not enforced; long-lived tokens become exploitable.",
                "Supabase Row-Level Security (RLS) policies misconfigured for certain roles.",
                "Third-party dependency with a known CVE included in the Docker image.",
            ],
            "mitigation": [
                "Enforce JWT + TOTP 2FA for all admin and manager roles.",
                "Run automated penetration tests quarterly; patch Critical/High CVEs within 7 days.",
                "Verify and test Supabase RLS policies for every table before each release.",
                "Conduct monthly dependency vulnerability scans (pip-audit / OWASP Dependency-Check).",
            ],
            "contingency": [
                "Immediately revoke all active sessions and force re-authentication on breach.",
                "Notify affected stakeholders and regulators within 72 hours per breach regulation.",
                "Engage a certified incident-response team to contain and remediate the breach.",
                "Perform a full post-incident security review and implement corrective controls.",
            ],
            "trigger": (
                "Audit logs show access patterns inconsistent with normal user behaviour, "
                "or a CVE with CVSS ≥ 8.0 is disclosed for a used dependency."
            ),
            "assigned_to": "Shivansh Srivastava  (Product Developer)",
            "originator": "Prahallad Padhan  (Product Owner)",
        },
        {
            "id": 4,
            "category": "DE",
            "prob": 50,
            "impact": 2,
            "description": (
                "Outage or degradation of critical third-party services — Supabase, Redis, "
                "Groq AI API — disrupts core IBMS operations."
            ),
            "subconditions": [
                "Single-region Supabase deployment is a single point of failure.",
                "Redis is not replicated; cache loss causes database overload.",
                "Groq AI API rate limits are exceeded during peak usage.",
            ],
            "mitigation": [
                "Implement in-memory fallback stores for all ERP ops when Supabase is unreachable.",
                "Deploy Redis with a replica and configure automatic failover.",
                "Cache Groq AI responses in Redis with appropriate TTL to reduce API call volume.",
                "Monitor /api/health endpoint; alert the team on any dependency degradation.",
            ],
            "contingency": [
                "Activate manual override procedures for critical business operations.",
                "Switch to backup Supabase project or PostgreSQL if primary is down > 1 hour.",
                "Disable AI-powered features temporarily; serve cached or default responses.",
                "Communicate service status and ETA to users via an in-app notification banner.",
            ],
            "trigger": (
                "Health-check endpoint reports a dependency as unhealthy for > 5 minutes, "
                "or error rate exceeds 5 % over a 10-minute window."
            ),
            "assigned_to": "Shivansh Srivastava  (Product Developer)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
        {
            "id": 5,
            "category": "ST",
            "prob": 45,
            "impact": 2,
            "description": (
                "Team members lack sufficient experience with FastAPI async patterns, "
                "the Frappe framework, Kubernetes, or ML model deployment, "
                "leading to low-quality code and rework."
            ),
            "subconditions": [
                "New members unfamiliar with Frappe's hook system and DocType conventions.",
                "Insufficient K8s experience leads to misconfigured deployment manifests.",
                "Lack of async Python expertise causes blocking I/O anti-patterns in FastAPI.",
            ],
            "mitigation": [
                "Conduct technology-specific onboarding workshops in the project's first two weeks.",
                "Pair junior developers with a senior mentor for all critical modules.",
                "Maintain an internal wiki with code patterns, pitfalls, and architectural decisions.",
                "Allocate a 15 % ramp-up time buffer in the first sprint per new team member.",
            ],
            "contingency": [
                "Engage specialist contractors for Frappe or K8s tasks if skill gap is confirmed.",
                "Reduce technical scope to areas within proven team competencies for v1.",
                "Obtain targeted training from certified providers (FastAPI, K8s, Frappe).",
                "Reassign modules to better-matched team members based on skill audit.",
            ],
            "trigger": (
                "Sprint velocity drops > 30 % below plan two sprints in a row, "
                "or code-review defect rate exceeds 40 % for a specific module."
            ),
            "assigned_to": "Ranveer Rai Khare  (Scrum Master)",
            "originator": "Prahallad Padhan  (Product Owner)",
        },
        {
            "id": 6,
            "category": "PS",
            "prob": 40,
            "impact": 2,
            "description": (
                "The IBMS platform degrades significantly under high concurrent user load, "
                "failing to serve real-time KPI WebSocket updates and ERP requests within SLA."
            ),
            "subconditions": [
                "WebSocket broadcast to many concurrent clients is not optimised.",
                "Supabase connection pool is exhausted under high concurrency (> 200 users).",
                "ML inference endpoints are not horizontally scaled and become bottlenecks.",
            ],
            "mitigation": [
                "Run load tests (Locust / k6) at target concurrency before each major release.",
                "Use Kubernetes HPA to auto-scale pods based on CPU/memory metrics.",
                "Cache KPI data in Redis with a 15-second TTL to reduce database load.",
                "Profile and optimise Supabase queries; add appropriate indexes.",
            ],
            "contingency": [
                "Throttle WebSocket clients and reduce KPI update frequency during peak load.",
                "Scale Kubernetes replicas manually beyond HPA limits if needed.",
                "Degrade gracefully: disable non-essential background jobs under high load.",
                "Negotiate a temporary SLA relaxation while performance improvements are applied.",
            ],
            "trigger": (
                "P95 API response time exceeds 2 seconds, or WebSocket disconnections "
                "exceed 10 % of active connections in any 5-minute window."
            ),
            "assigned_to": "Shivansh Srivastava  (Product Developer)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
        {
            "id": 7,
            "category": "CU",
            "prob": 35,
            "impact": 3,
            "description": (
                "Enterprise clients resist adopting IBMS due to unfamiliarity with the new system, "
                "fear of data-migration risks, or preference for legacy tools."
            ),
            "subconditions": [
                "End users were not involved in the requirements-gathering phase.",
                "No structured change-management or training programme was planned.",
                "Data migration from legacy systems is perceived as high risk by the IT department.",
            ],
            "mitigation": [
                "Involve key end-user representatives in sprint demos and UAT sessions.",
                "Develop comprehensive user guides and video tutorials before go-live.",
                "Offer a phased rollout with a parallel-run period to build user confidence.",
                "Assign a dedicated Customer Success point-of-contact for the first 3 months.",
            ],
            "contingency": [
                "Extend the parallel-run period and delay full cutover until adoption metrics are met.",
                "Provide on-site training sessions for power users at no extra charge.",
                "Simplify the UI for less tech-savvy users based on usability testing feedback.",
                "Set up a dedicated support channel (email / Slack) during the adoption phase.",
            ],
            "trigger": (
                "User adoption rate is below 60 % after 4 weeks of go-live, "
                "or NPS score is below 20 in the first post-launch survey."
            ),
            "assigned_to": "Prahallad Padhan  (Product Owner)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
        {
            "id": 8,
            "category": "BU",
            "prob": 30,
            "impact": 2,
            "description": (
                "The project budget is exceeded due to underestimation of cloud infrastructure "
                "costs (GCP / AWS), AI API usage fees (Groq), and third-party service subscriptions."
            ),
            "subconditions": [
                "Cloud resource costs were estimated at flat rates without usage-spike modelling.",
                "Groq AI API charges are based on token volume, which was underestimated.",
                "Supabase costs scale with database size and API requests beyond the free tier.",
            ],
            "mitigation": [
                "Build a detailed cloud cost model using GCP/AWS calculators before each sprint.",
                "Set billing alerts at 70 % and 90 % thresholds on all cloud accounts.",
                "Review infrastructure costs in bi-weekly budget review meetings.",
                "Maintain a 15 % contingency reserve in the project budget.",
            ],
            "contingency": [
                "Right-size cloud resources (reduce instance types, shut unused regions) immediately.",
                "Replace Groq AI with a self-hosted open-source model (Ollama / Llama) if costs spike.",
                "Pause non-essential staging environments during cost-overrun periods.",
                "Initiate a formal budget re-baseline request with the project sponsor.",
            ],
            "trigger": (
                "Monthly cloud spend exceeds the monthly budget by > 15 %, "
                "or cumulative project spend reaches 85 % of the total approved budget "
                "before 75 % of milestones are complete."
            ),
            "assigned_to": "Prahallad Padhan  (Product Owner)",
            "originator": "Ranveer Rai Khare  (Scrum Master)",
        },
    ]

    ris_order = sorted(RIS_DATA, key=lambda r: r["prob"], reverse=True)

    LABEL_W  = Cm(3.8)
    CONTENT_W = Cm(12.12)   # 3.8 + 12.12 = 15.92 cm

    def ris_simple_row(tbl, label, content, label_bg=MID_BLUE, content_bg=WHITE, sz=9.5):
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        lc.width = LABEL_W;  vc.width = CONTENT_W
        set_cell_bg(lc, label_bg); cell_para(lc, label, bold=True, color=WHITE, size=9.5)
        set_cell_bg(vc, content_bg); cell_para(vc, content, size=sz)

    def ris_list_row(tbl, label, items, label_bg=MID_BLUE, content_bg=WHITE):
        row = tbl.add_row()
        lc, vc = row.cells[0], row.cells[1]
        lc.width = LABEL_W;  vc.width = CONTENT_W
        set_cell_bg(lc, label_bg); cell_para(lc, label, bold=True, color=WHITE, size=9.5)
        set_cell_bg(vc, content_bg)
        vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        for i, item in enumerate(items):
            p = vc.paragraphs[0] if i == 0 else vc.add_paragraph()
            p.clear()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            rn = p.add_run(f"  {i + 1}.  {item}")
            rn.font.size      = Pt(9.5)
            rn.font.color.rgb = BLACK

    for idx, ris in enumerate(ris_order):
        if idx > 0:
            doc.add_page_break()

        # ── Banner ───────────────────────────────────────────────────────
        banner_tbl = doc.add_table(rows=1, cols=1)
        banner_tbl.style = "Table Grid"
        set_table_fullwidth(banner_tbl)
        bc = banner_tbl.rows[0].cells[0]
        set_cell_bg(bc, DARK_BLUE)
        cell_para(
            bc,
            f"Risk Information Sheet (RIS)   |   Risk ID: R{ris['id']}   |   "
            f"Probability: {ris['prob']} % ({prob_label(ris['prob'])})   |   "
            f"Impact: {ris['impact']} – {IMPACT_LABELS[ris['impact']]}",
            bold=True, color=WHITE, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=3, space_after=3,
        )

        # ── Meta strip (one row, 4 cols) ─────────────────────────────────
        meta_tbl = doc.add_table(rows=1, cols=4)
        meta_tbl.style = "Table Grid"
        set_table_fullwidth(meta_tbl)
        for ci in range(4):
            meta_tbl.columns[ci].width = Cm(3.98)
        meta_items = [
            ("Date",        "24 / 04 / 2026"),
            ("Category",    f"{ris['category']} – {CATEGORY_FULL[ris['category']]}"),
            ("Assigned To", ris["assigned_to"]),
            ("Originator",  ris["originator"]),
        ]
        for ci, (lbl, val) in enumerate(meta_items):
            mc = meta_tbl.rows[0].cells[ci]
            set_cell_bg(mc, LIGHT_BLU)
            p = mc.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r1 = p.add_run(lbl + ":  ")
            r1.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = DARK_BLUE
            r2 = p.add_run(val)
            r2.font.size = Pt(8.5); r2.font.color.rgb = BLACK

        # ── Body table ───────────────────────────────────────────────────
        body_tbl = doc.add_table(rows=0, cols=2)
        set_tbl_style(body_tbl)
        body_tbl.columns[0].width = LABEL_W
        body_tbl.columns[1].width = CONTENT_W

        ris_simple_row(body_tbl, "Description",
                       ris["description"], content_bg=LIGHT_GRN, sz=9.5)

        sub_text = "\n".join(
            f"  Sub-condition {i+1}:  {s}" for i, s in enumerate(ris["subconditions"])
        )
        ris_simple_row(body_tbl, "Refinement &\nContext\n(Sub-conditions)",
                       sub_text, content_bg=WHITE, sz=9.5)

        ris_list_row(body_tbl, "Mitigation &\nMonitoring\nStrategies",
                     ris["mitigation"], content_bg=LIGHT_GRN)

        ris_list_row(body_tbl, "Contingency Plan\n& Management",
                     ris["contingency"], content_bg=WHITE)

        ris_simple_row(body_tbl, "Trigger",
                       ris["trigger"], content_bg=LIGHT_GRN, sz=9.5)

        keep_table_on_page(body_tbl)

    # ── SECTION 5 — SUMMARY (no new page, appended after last RIS) ───────────
    sp = doc.add_paragraph()
    no_space_after(sp)

    add_heading(doc, "5.  Summary & Review Notes")
    add_body(doc,
        "Each identified risk has an individual Risk Information Sheet (RIS) as documented above.  "
        "This RMMM Plan must be reviewed at the start of every sprint and updated whenever a "
        "risk's probability or impact changes.  New risks discovered during the project lifecycle "
        "must be added immediately with a new Risk ID and a corresponding RIS.")
    add_body(doc,
        "Probability values are consensus estimates agreed by the team using the Wideband Delphi "
        "technique.  Impact levels follow the four-tier severity scale (Catastrophic → Negligible) "
        "as defined in Pressman's Software Engineering textbook.", italic=True)

    # ── SAVE ─────────────────────────────────────────────────────────────────
    out_path = r"d:\Integrated-Management-Business-Suite\docs\RMMM_Plan_IBMS_v2.docx"
    doc.save(out_path)
    print(f"[OK] Document saved → {out_path}")


if __name__ == "__main__":
    build_document()

